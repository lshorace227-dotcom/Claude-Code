"""
亞朵生活控股 (ATAT) — 2026 Q1 業績更新報告
產出：ATAT_Q1_2026_Earnings_Update.docx
機構研究格式，10-12 頁，內嵌 10 張圖表
報告語言：繁體中文（專業術語保留英文）
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHARTS = os.path.join(os.path.dirname(__file__), 'charts')
OUT_FILE = os.path.join(os.path.dirname(__file__), 'ATAT_Q1_2026_Earnings_Update.docx')

# ─── 色彩 ────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x4F, 0x8A)
ORANGE = RGBColor(0xE8, 0x72, 0x2A)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREEN  = RGBColor(0x2E, 0x8B, 0x57)
GRAY   = RGBColor(0x7F, 0x8C, 0x8D)
LGRAY  = RGBColor(0xF2, 0xF3, 0xF4)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
GOLD   = RGBColor(0xD4, 0xAC, 0x0D)

ASCII_FONT = 'Times New Roman'
CJK_FONT   = 'PMingLiU'  # 繁體中文通用字體


def _apply_cjk_font(run, ascii_font=ASCII_FONT, cjk_font=CJK_FONT):
    """確保中文字符使用 East Asian 字體。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), cjk_font)
    rFonts.set(qn('w:cs'), ascii_font)


def set_cell_bg(cell, rgb_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)


def para_run(para, text, bold=False, italic=False, size=10, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = ASCII_FONT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    _apply_cjk_font(run)
    return run


def add_heading(doc, text, level=1, color=NAVY, size=None):
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = ASCII_FONT
    run.font.size = Pt(size or sizes.get(level, 11))
    run.font.color.rgb = color
    _apply_cjk_font(run)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '1B4F8A')
        pBdr.append(bottom)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, indent=False, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = ASCII_FONT
    run.font.size = Pt(10)
    _apply_cjk_font(run)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullet_takeaway(doc, bold_prefix, body_text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run('■  ')
    r1.font.color.rgb = color
    r1.font.size = Pt(10)
    r1.font.name = ASCII_FONT
    _apply_cjk_font(r1)
    r2 = p.add_run(bold_prefix + '：')
    r2.bold = True
    r2.font.name = ASCII_FONT
    r2.font.size = Pt(10)
    r2.font.color.rgb = color
    _apply_cjk_font(r2)
    r3 = p.add_run(body_text)
    r3.font.name = ASCII_FONT
    r3.font.size = Pt(10)
    _apply_cjk_font(r3)
    return p


def add_chart(doc, filename, width=6.0, caption=None):
    path = os.path.join(CHARTS, filename)
    if not os.path.exists(path):
        print(f"  WARNING: 找不到圖檔：{path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.name = ASCII_FONT
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        r.italic = True
        _apply_cjk_font(r)
        cap.paragraph_format.space_after = Pt(6)


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.0):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def divider(doc, color='1B4F8A'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '1B4F8A')
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(color_el)
    rPr.append(u_el)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ASCII_FONT)
    rFonts.set(qn('w:hAnsi'), ASCII_FONT)
    rFonts.set(qn('w:eastAsia'), CJK_FONT)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)
    return hyperlink


# ═══════════════════════════════════════════════════════════════════
# 報告主體
# ═══════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_page_margins(doc)

    normal = doc.styles['Normal']
    normal.font.name = ASCII_FONT
    normal.font.size = Pt(10)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), ASCII_FONT)
    rFonts.set(qn('w:hAnsi'), ASCII_FONT)
    rFonts.set(qn('w:eastAsia'), CJK_FONT)

    # ═══════════════════════════════════════════════════════════════
    # 第 1 頁：封面與摘要
    # ═══════════════════════════════════════════════════════════════
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '1B4F8A')

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('業績更新 | EARNINGS UPDATE')
    r.font.name = ASCII_FONT; r.font.size = Pt(9); r.bold = True
    r.font.color.rgb = WHITE; _apply_cjk_font(r)

    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('亞朵生活控股 Atour Lifestyle Holdings Ltd. (NASDAQ: ATAT)')
    r2.font.name = ASCII_FONT; r2.font.size = Pt(16); r2.bold = True
    r2.font.color.rgb = WHITE; _apply_cjk_font(r2)

    p3 = cell.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('2026 年第一季度業績 | 報告日期：2026 年 5 月 27 日')
    r3.font.name = ASCII_FONT; r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1); r3.italic = True
    _apply_cjk_font(r3)

    cell.paragraphs[0].paragraph_format.space_before = Pt(8)
    for pp in cell.paragraphs:
        pp.paragraph_format.space_after = Pt(4)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 評級 / 目標股價方塊 ──
    rt = doc.add_table(rows=1, cols=4)
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ['投資評級', '12 個月目標股價', '前次目標股價', '現價']
    values = ['買入 / BUY', 'US$63', 'US$58', 'US$36.50']
    colors_bg = ['1B4F8A', 'E8722A', '7F8C8D', '2C3E50']
    for cell_x, lbl, val, bg in zip(rt.row_cells(0), labels, values, colors_bg):
        set_cell_bg(cell_x, bg)
        p1 = cell_x.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p1.add_run(lbl)
        rl.font.name = ASCII_FONT; rl.font.size = Pt(8)
        rl.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1)
        _apply_cjk_font(rl)
        p2c = cell_x.add_paragraph(); p2c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = p2c.add_run(val)
        rv.font.name = ASCII_FONT; rv.font.size = Pt(13); rv.bold = True
        rv.font.color.rgb = WHITE
        _apply_cjk_font(rv)
        p1.paragraph_format.space_before = Pt(6)
        p2c.paragraph_format.space_after = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 關鍵摘要 ──
    add_heading(doc, '關鍵投資要點 — Q1 2026 業績全面超預期', level=1)

    takeaways = [
        ('營收大幅超預期 (+10.2%)', '2026 Q1 淨營收人民幣 28.11 億元（美金 4.14 億元），年增 +47.5%；高於 Bloomberg 一致預期美金 3.756 億元約 2,050 萬美元或 +10.2%。零售與加盟管理酒店業務雙引擎驅動營收創歷史單季新高。', GREEN),
        ('EPS 大幅超預期 (+40.5%)', '經調整每 ADS 盈利 US$0.52，較市場一致預期 US$0.37 高出 US$0.15 或 +40.5%。經調整淨利潤年增 +42.0% 至人民幣 4.90 億元；GAAP 淨利潤年增 +90.3% 至人民幣 4.63 億元。', GREEN),
        ('零售業務指引上調', '管理層將 FY2026 零售業務營收增速指引由原 25-30% 上調至 30-35%，反映「深睡科技恆溫被 Pro 3.0」等新品在 45 天內 GMV 突破人民幣 1 億元的強勁市場反應。FY2026 總營收指引維持 +24-28% 不變。', NAVY),
        ('酒店網絡擴張穩健', 'Q1 季度新開業 110 家酒店、關閉 37 家、淨增加 73 家，使在營酒店數達到 2,088 家（年增 +20.9%），會員數突破 1.16 億人（年增 +20%）。在建管道 751 家為後續成長提供堅實支撐。', NAVY),
        ('成熟酒店仍承壓', '整體 RevPAR 為人民幣 312 元（年增 +2.4%），ADR 達人民幣 427 元（年增 +2.1%），入住率 70.6%；惟營運超過 18 個月的成熟酒店 RevPAR 同比僅達 98.3%，反映行業階段性復甦的非線性特徵。', ORANGE),
        ('資本配置友善：股息維持高派息', '董事會宣派每股普通股 US$0.18（每 ADS US$0.54）股息，總額約 US$7,200 萬元，相當於 FY2025 淨利潤約 31%；公司累計回購已超過 US$1 億元，承諾接近 100% 派息政策。', NAVY),
        ('維持買入評級 | 目標股價上調 US$58 → US$63', '基於 FY2027E 經調整 EPS US$2.85 × 22x P/E（同業均值約 23x），上調目標股價至 US$63，較現價 US$36.50 隱含上行空間 +72%。下行情境（零售放緩 + RevPAR 走弱）下行支撐 US$39；上行情境（估值重估）有望達 US$87。', GREEN),
    ]
    for prefix, body, color in takeaways:
        add_bullet_takeaway(doc, prefix, body, color)

    # 圖 1：季度營收
    add_chart(doc, 'chart_01_quarterly_revenue.png', width=6.4)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 第 2-3 頁：業績詳細解讀
    # ═══════════════════════════════════════════════════════════════
    add_heading(doc, '一、業績解讀：營收與獲利雙超預期', level=1)

    add_body(doc,
        '亞朵生活控股於 2026 年 5 月 13 日（美國東部時間盤前）公佈 2026 年第一季度業績，'
        '核心財務與營運指標普遍優於市場一致預期。本季淨營收人民幣 28.11 億元（美金 4.14 億元），'
        '年增 +47.5%，較 Bloomberg 一致預期 3.756 億美元超預期 +10.2%；'
        'GAAP 淨利潤人民幣 4.63 億元（年增 +90.3%），經調整淨利潤人民幣 4.90 億元（年增 +42.0%）。'
        '經調整 EBITDA 達人民幣 7.16 億元，年增 +51.1%，相應利潤率擴張 60bps 至 25.5%。'
        '管理層在隨後召開的業績說明會上強調，季度業績反映了「酒店網絡擴張 + 零售品牌力放大」'
        '的雙引擎模式持續見效，並基於此將 FY2026 零售業務指引由原 +25-30% 上調至 +30-35%。')

    add_heading(doc, '1.1 分部營收結構：零售與加盟酒店共同驅動', level=2)

    add_body(doc,
        '從營收結構看，本季增長動能呈現「兩條腿走路」格局：（1）加盟管理酒店營收年增 +51.9% '
        '至人民幣 15.68 億元，佔總營收 55.8%，反映酒店家數擴張 +20.9% 帶動的規模效應，加上'
        '管理費率穩中有升及供應鏈業務貢獻；（2）零售業務營收年增 +54.4% 至人民幣 10.71 億元，'
        '佔總營收 38.1%，較去年同期 36.4% 進一步提升，「深睡」系列床品的爆款效應持續放大。'
        '相對而言，租賃及自營酒店營收年減 -8.0% 至人民幣 1.18 億元，反映公司持續優化租賃酒店組合、'
        '將部分自營門店轉為加盟模式的策略選擇。其他業務營收年增 +4.7% 至人民幣 5,400 萬元，'
        '結構相對穩定。')

    add_chart(doc, 'chart_02_segment_revenue.png', width=6.5)

    add_heading(doc, '1.2 獲利能力：GAAP 與 Non-GAAP 出現分化', level=2)

    add_body(doc,
        'GAAP 淨利率由去年同期 12.8% 大幅提升至 16.5%（+370bps），主要受惠於零售業務規模化帶來的'
        '經營槓桿釋放，以及股權激勵費用較去年同期下降的基數效應。然而，更具經營意義的經調整淨利率'
        '由 18.1% 微降 70bps 至 17.4%，反映公司在零售品牌投入、新店爬坡期費用、員工成本'
        '及行銷費用等方面的前置投入。經調整 EBITDA 利潤率為 25.5%，較去年同期 24.9% 略升 60bps，'
        '驗證底層獲利能力仍維持健康擴張軌跡。我們認為，短期經調整淨利率輕度收縮屬於'
        '「成長性投入換取長期市場份額」的合理代價，並不構成投資邏輯破壞。')

    add_chart(doc, 'chart_03_profitability.png', width=6.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 第 4-5 頁：營運指標與酒店網絡
    # ═══════════════════════════════════════════════════════════════
    add_heading(doc, '二、核心營運指標解讀', level=1)

    add_heading(doc, '2.1 RevPAR / ADR / 入住率：整體穩中有升，成熟店仍承壓', level=2)

    add_body(doc,
        '本季整體酒店組合 RevPAR 為人民幣 311.6 元，為去年同期 102.4%，ADR 人民幣 427 元（YoY +2.1%），'
        '入住率 70.6%（去年同期 70.2%）。Q1 屬於酒店業傳統淡季，與 Q4 2025 旺季 RevPAR 人民幣 336 元、'
        '入住率 76.1% 相比出現季節性回落屬正常。'
        '但值得關注的是，營運超過 18 個月的成熟酒店組合 RevPAR 為去年同期 98.3%（入住率 99.2%、'
        'ADR 99.4%），表明在中國酒店業整體進入「波動性復甦階段」的宏觀背景下，'
        '同店端的價量平衡仍處於微幅承壓狀態。管理層在業績會上以「波動性復甦」（fluctuating recovery）'
        '形容當前需求環境，並坦承 RevPAR 趨勢「在短期會出現波動」。')

    add_body(doc,
        '我們認為，整體 RevPAR 同比 +2.4% 的表現好於同業（華住集團、錦江、首旅 Q1 同店 RevPAR '
        '預期均處於同比微降區間），凸顯了亞朵中高端定位下品牌溢價的相對韌性。'
        '但成熟店 RevPAR 同比 -1.7% 的數據仍需後續觀察 Q2 起的同店趨勢是否能扭轉，'
        '這將是後續股價催化劑的關鍵指標之一。')

    add_chart(doc, 'chart_04_revpar_trends.png', width=6.5)

    add_heading(doc, '2.2 酒店網絡擴張：管道蓄水充沛', level=2)

    add_body(doc,
        '截至 2026 年 3 月 31 日，亞朵在營酒店達 2,088 家（年增 +20.9%），客房總數達 232,298 間'
        '（年增 +19.4%）。其中 Q1 新開業 110 家、關閉 37 家（含主動優化的低效門店及自然退出'
        '的少量門店），淨增加 73 家。重要的是，公司在建管道達 751 家，較 2025 年底進一步擴大，'
        '為下半年集中開業提供堅實儲備。管理層在電話會議中明確表示，全年新增門店目標維持不變，'
        '隱含 Q2-Q4 將維持高強度開業節奏。')

    add_body(doc,
        '會員端，註冊個人會員人數達 1.16 億人（年增 +20%），與酒店家數擴張節奏匹配。'
        '會員體系是亞朵差異化競爭優勢的核心：高黏性會員不僅貢獻直訂佔比超過行業平均，'
        '更是零售業務復購率的關鍵基礎。我們估算，目前每名活躍會員年均零售消費約人民幣 35-40 元，'
        '隨著零售品類擴張及複購提升，這一指標仍有 1.5-2 倍的提升空間。')

    add_chart(doc, 'chart_05_hotel_network.png', width=6.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 第 6-7 頁：超預期分析與成本費用
    # ═══════════════════════════════════════════════════════════════
    add_heading(doc, '三、業績超預期分析（Beat / Miss）', level=1)

    add_body(doc,
        '本季業績相對 Bloomberg 一致預期的超預期幅度，無論在營收端還是利潤端均屬於「實質超預期」'
        '（meaningful beat）級別：（1）營收 4.14 億美元，較預期 3.756 億美元超預期 +10.2%；'
        '（2）EPS US$0.52 較預期 US$0.37 超預期 +40.5%；（3）經調整 EBITDA 約 1.04 億美元，'
        '較預期 0.92 億美元超預期 +13.0%；（4）GAAP 淨利潤 0.67 億美元，較預期 0.49 億美元'
        '超預期 +36.7%。超預期主要來自三大因素：（a）零售業務的爆款驅動效應遠超賣方建模'
        '的線性外推；（b）加盟管理費率受惠於組合升級（部分新店為中端偏上品牌）；'
        '（c）股權激勵費用低於市場預期的常規假設。')

    add_chart(doc, 'chart_06_beat_miss.png', width=6.5)

    add_heading(doc, '四、營運費用結構分析', level=1)

    add_body(doc,
        '本季營運費用呈現「投入優先、結構分化」的特徵。酒店營運成本年增 +38.5% 至人民幣 11.36 億元，'
        '增速顯著低於酒店相關營收（管理+租賃）合計增速 +45.8%，反映酒店業務的營運槓桿正在顯現。'
        '零售商品成本年增 +45.7% 至人民幣 5.07 億元，低於零售營收增速 +54.4%，'
        '相應零售毛利率擴張至 52.6%（去年同期 49.9%），表明零售業務在規模化過程中議價能力'
        '及供應鏈效率均有改善。')

    add_body(doc,
        '另一方面，銷售與行銷費用年增 +93.7% 至人民幣 4.01 億元，明顯高於營收增速，'
        '反映公司在零售品牌建設、直播電商、社交媒體種草等方面的密集投入。'
        '一般及行政費用年增 +27.3%、研發費用年增 +38.9%，均處於相對溫和水平。'
        '我們對此的解讀是：管理層在零售業務的「黃金窗口期」選擇加大行銷投入，'
        '優先換取市場份額與品牌心智，這在短期會壓縮利潤率，但有助於長期建立護城河。')

    add_chart(doc, 'chart_07_opex_structure.png', width=6.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 第 8 頁：指引與估算更新
    # ═══════════════════════════════════════════════════════════════
    add_heading(doc, '五、管理層指引與估算更新', level=1)

    add_heading(doc, '5.1 指引變化：零售業務上調 +5pp', level=2)

    add_body(doc,
        '管理層在業績說明會上做出兩項關鍵指引調整：（1）將 FY2026 零售業務營收增速指引從原 '
        '+25-30% 上調至 +30-35%（+5pp），主要基於 Q1 強勁表現及「深睡科技恆溫被 Pro 3.0」、'
        '新一代旅行箱等新品的市場驗證；（2）FY2026 總營收增速指引維持 +24-28% 不變，'
        '隱含公司對酒店端 RevPAR 趨勢採取相對審慎的假設。')

    add_body(doc,
        '基於指引上調以及 Q1 業績全面超預期，我們將 FY2026E 營收預估從人民幣 98 億元上調 +6.1% '
        '至人民幣 104 億元；經調整淨利潤預估由人民幣 17.5 億元上調 +6.3% 至人民幣 18.6 億元；'
        '經調整 EPS 預估由 US$2.10 上調 +8.6% 至 US$2.28。FY2027E 經調整 EPS 上調至 US$2.85，'
        '較原預估上修 +9.6%，反映零售業務複合成長加速。')

    add_chart(doc, 'chart_08_guidance_revisions.png', width=6.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 第 9 頁：估值對比
    # ═══════════════════════════════════════════════════════════════
    add_heading(doc, '六、估值對比與目標股價', level=1)

    add_heading(doc, '6.1 同業估值對比：成長性領先但估值折讓', level=2)

    add_body(doc,
        '基於 Bloomberg / FactSet 一致預期，亞朵 FY2026E P/E 約 16.5x，顯著低於同業均值 23.3x。'
        '即便與中國酒店同業相比，亞朵估值亦處於折讓區間：華住集團 21.3x、錦江酒店 25.8x、'
        '首旅酒店 28.0x；與國際同業 IHG 22.4x、萬豪 25.6x 對比，折讓更為明顯。'
        '更值得關注的是，亞朵 FY2026E 營收增速約 26%，為同業中最高，相應 PEG 僅 0.6x，'
        '在「高成長 + 低估值」維度具備突出吸引力。')

    add_body(doc,
        '我們認為亞朵估值折讓的主要原因有三：（1）中概股流動性折讓持續存在；'
        '（2）零售業務佔比上升使市場對其行業歸類存在爭議（介於酒店、消費品牌與電商之間）；'
        '（3）部分投資者對中國酒店業階段性復甦的可持續性存在疑慮。'
        '我們認為，隨著零售業務持續驗證爆品邏輯，估值折讓有望逐步收斂。')

    add_chart(doc, 'chart_09_peer_valuation.png', width=6.5)

    add_heading(doc, '6.2 目標股價情境分析', level=2)

    add_body(doc,
        '基於 FY2027E 經調整 EPS US$2.85 與 22x 目標 P/E（接近同業均值小幅折讓），'
        '我們將 12 個月目標股價由 US$58 上調至 US$63，較 2026/05/26 收盤價 US$36.50 隱含'
        '上行空間 +72%。情境分析：（1）熊市情境（零售增速放緩至 +20%、RevPAR 同店持續同比 -5%）'
        '對應目標股價 US$39，下行空間 +7%；（2）牛市情境（零售加速至 +40%、估值重估至 26x P/E）'
        '對應目標股價 US$87，上行空間 +138%。風險回報比約為 1:4.7，具備明顯不對稱優勢。')

    add_chart(doc, 'chart_10_price_target.png', width=6.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 第 10 頁：估算彙總表 / 投資論點更新 / 風險
    # ═══════════════════════════════════════════════════════════════
    add_heading(doc, '七、預估彙總（FY2024A — FY2027E）', level=1)

    # 預估表
    est_tbl = doc.add_table(rows=10, cols=6)
    est_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    est_tbl.autofit = False

    headers = ['指標（RMB M 除註明）', 'FY2024A', 'FY2025A', 'FY2026E\n(新)', 'FY2026E\n(舊)', 'FY2027E']
    data = [
        ['淨營收',                '6,411',  '8,168',  '10,400', '9,800',  '12,950'],
        ['  年增率 (%)',          '+55.3%', '+27.4%', '+27.3%', '+20.0%', '+24.5%'],
        ['毛利',                  '2,800',  '3,580',  '4,650',  '4,360',  '5,860'],
        ['  毛利率 (%)',          '43.7%',  '43.8%',  '44.7%',  '44.5%',  '45.2%'],
        ['經調整 EBITDA',         '1,420',  '1,920',  '2,580',  '2,400',  '3,290'],
        ['  EBITDA 利潤率 (%)',   '22.2%',  '23.5%',  '24.8%',  '24.5%',  '25.4%'],
        ['經調整淨利潤',          '1,260',  '1,420',  '1,860',  '1,750',  '2,330'],
        ['經調整 EPS (US$/ADS)',  '$1.42',  '$1.66',  '$2.28',  '$2.10',  '$2.85'],
        ['FCF',                   '1,100',  '1,550',  '2,100',  '1,950',  '2,800'],
    ]

    cells = est_tbl.row_cells(0)
    for i, h in enumerate(headers):
        set_cell_bg(cells[i], '1B4F8A')
        cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = ASCII_FONT; r.font.size = Pt(9); r.bold = True
        r.font.color.rgb = WHITE; _apply_cjk_font(r)

    for r_idx, row_data in enumerate(data, start=1):
        cells = est_tbl.row_cells(r_idx)
        if r_idx % 2 == 0:
            for c in cells:
                set_cell_bg(c, 'F8F9FA')
        for c_idx, val in enumerate(row_data):
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(val)
            run.font.name = ASCII_FONT; run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = val.startswith(('淨營收', '毛利', '經調整', 'FCF'))
            elif c_idx == 3 and r_idx in [1, 7]:  # 新預估數重點標示
                run.bold = True; run.font.color.rgb = ORANGE
            _apply_cjk_font(run)

    # caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('資料來源：亞朵生活控股業績公告；本研究團隊預估；2026/05/26')
    r.font.name = ASCII_FONT; r.font.size = Pt(8); r.font.color.rgb = GRAY; r.italic = True
    _apply_cjk_font(r)
    p.paragraph_format.space_after = Pt(8)

    # ── 投資論點更新 ──
    add_heading(doc, '八、投資論點更新', level=1)

    add_heading(doc, '8.1 維持的核心論點', level=2)
    for prefix, body in [
        ('酒店網絡擴張仍處於高速軌道',
         '在建管道 751 家為未來 8-10 個季度的開業節奏提供強勁支撐，公司「中高端 + 加盟為主」的'
         '商業模式在資本效率與品牌可控性間取得了良好平衡。'),
        ('零售業務驗證「品牌 + 場景」邏輯',
         '「深睡」系列恆溫被在 45 天 GMV 突破 1 億元，再次證明亞朵已建立「酒店體驗 → 會員信任 → '
         '居家品牌」的轉化閉環。此模式在全球酒店業中具備獨特性，且具備複製到更多品類的潛力。'),
        ('資本配置紀律性強', '~100% 派息政策疊加累計超 1 億美元的回購，反映管理層對股東回報的'
         '高度重視；同時 RMB 3.7B 現金、僅 RMB 242M 負債的健康資產負債表為未來戰略選擇'
         '留出充足彈性。'),
    ]:
        add_bullet_takeaway(doc, prefix, body)

    add_heading(doc, '8.2 修正後的論點', level=2)
    for prefix, body in [
        ('零售業務貢獻時程提前',
         '原預期零售營收佔比在 FY2027 達到 40%，現基於指引上調及產品矩陣擴張節奏，'
         '預期 FY2026 末即可達到 40-42% 區間，意味著估值體系應更多參考消費品牌而非'
         '純酒店業者的估值錨。'),
        ('成熟酒店 RevPAR 短期難回正成長',
         '基於 Q1 同店 RevPAR -1.7% 的表現及管理層「波動性復甦」的定調，我們將 FY2026 '
         '同店 RevPAR 假設由原 +1.5% 下修至 -1.0%。但因新店貢獻持續放大，整體 RevPAR 仍能維持'
         '微幅同比上升。'),
    ]:
        add_bullet_takeaway(doc, prefix, body, ORANGE)

    add_heading(doc, '8.3 主要風險', level=2)
    for prefix, body in [
        ('行業需求復甦不及預期',
         '若中國消費信心反覆，商務出行與休閒旅遊需求復甦延後，將直接壓縮酒店端 RevPAR。'),
        ('零售業務爆品依賴度高',
         '當前零售業務增長高度依賴「深睡」系列床品爆款，若新品迭代節奏放緩或競品快速跟進，'
         '可能影響高速成長軌道的延續性。'),
        ('競爭加劇',
         '華住、錦江等同業正加快中高端品牌矩陣擴張，可能在優質物業獲取與管理費率方面對亞朵形成壓力。'),
        ('中概股政策與流動性風險',
         '美中監管政策、ADR 退市風險與 VIE 結構風險仍是長期需要考量的系統性因素。'),
    ]:
        add_bullet_takeaway(doc, prefix, body, RED)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 第 11 頁：資料來源與引用
    # ═══════════════════════════════════════════════════════════════
    add_heading(doc, '九、資料來源與引用', level=1)

    add_heading(doc, '9.1 主要公司資料', level=3)

    # Earnings release
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• 2026 年第一季度業績公告（2026 年 5 月 13 日）：')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)
    add_hyperlink(p, 'SEC Form 6-K — Atour Lifestyle Q1 2026 Earnings Release',
                  'https://www.sec.gov/Archives/edgar/data/0001853717/000110465926050659/tm2612813d1_ex99-1.htm')

    # Earnings call transcript
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• 業績說明會逐字稿（2026 年 5 月 13 日）：')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)
    add_hyperlink(p, 'Q1 2026 Earnings Call Transcript',
                  'https://www.sahmcapital.com/news/content/atour-lifestyle-holdings-q1-2026-earnings-call-transcript-2026-05-13')

    # Investor presentation
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• 業績電話會議簡報（2026 年 5 月 15 日）：')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)
    add_hyperlink(p, 'Atour Lifestyle Q1 2026 Earnings Call Presentation',
                  'https://seekingalpha.com/article/4905511-atour-lifestyle-holdings-limited-2026-q1-results-earnings-call-presentation')

    # SEC filings index
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• SEC EDGAR 申報文件總覽（Atour 6-K 系列）：')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)
    add_hyperlink(p, 'SEC EDGAR — Atour Lifestyle Holdings Form 6-K Filings',
                  'https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001853717&type=6-K&dateb=&owner=include&count=40')

    add_heading(doc, '9.2 一致預期與市場資料', level=3)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• Bloomberg / FactSet 一致預期數據（2026 年 5 月 26 日，業績公告後更新）')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• 股價來源：NASDAQ 收盤價（US$36.50，2026 年 5 月 26 日）：')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)
    add_hyperlink(p, 'Atour Lifestyle Holdings (ATAT) — NASDAQ',
                  'https://www.nasdaq.com/market-activity/stocks/atat')

    add_heading(doc, '9.3 第三方分析與評論', level=3)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• StockTitan 業績綜述：')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)
    add_hyperlink(p, 'Atour Lifestyle Q1 Net Income Jumps 90% to RMB463m',
                  'https://www.stocktitan.net/news/ATAT/atour-lifestyle-holdings-limited-reports-first-quarter-2026-8o1nyaxkhbgb.html')

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• The Globe and Mail 業績會議重點：')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)
    add_hyperlink(p, 'Atour Lifestyle\'s Earnings Call Highlights Growth And Caution',
                  'https://www.theglobeandmail.com/investing/markets/stocks/ATAT-Q/pressreleases/2056683/atour-lifestyles-earnings-call-highlights-growth-and-caution/')

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• Investing.com 業績會議逐字稿摘要：')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)
    add_hyperlink(p, 'Atour Lifestyle Holdings Sees 47.5% Revenue Growth in Q1 2026',
                  'https://www.investing.com/news/transcripts/earnings-call-transcript-atour-lifestyle-holdings-sees-475-revenue-growth-in-q1-2026-93CH-4689258')

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run('• Quartr — 公司投資人關係資料彙整：')
    r.font.name = ASCII_FONT; r.font.size = Pt(10); _apply_cjk_font(r)
    add_hyperlink(p, 'Atour Lifestyle (ATAT) Investor Relations, Earnings Summary & Outlook',
                  'https://quartr.com/companies/atour-lifestyle-holdings-limited_16849')

    # ── 免責聲明 ──
    divider(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run('免責聲明：')
    r.font.name = ASCII_FONT; r.font.size = Pt(8); r.bold = True
    r.font.color.rgb = GRAY; _apply_cjk_font(r)
    r2 = p.add_run(
        '本報告基於 2026 年 5 月 13 日亞朵生活控股公開披露之 Q1 2026 業績資料及隨後的業績說明會逐字稿撰寫，'
        '所有數據截至 2026 年 5 月 27 日。本報告所載之預估、評級、目標股價、情境分析均為'
        '基於當前可得資訊的判斷，未來可能因市場條件、公司基本面變動或宏觀環境變化而調整。'
        '報告中所引用之 Bloomberg / FactSet 一致預期、同業估值數據及股價來源已於正文及附錄中標註。'
        '本報告不構成任何投資建議或要約，投資人應自行評估並承擔投資風險。'
        '請特別注意 ADR 流動性、VIE 結構及美中監管政策等中概股特有的系統性風險。')
    r2.font.name = ASCII_FONT; r2.font.size = Pt(8); r2.font.color.rgb = GRAY
    r2.italic = True; _apply_cjk_font(r2)

    # ── Save ──
    doc.save(OUT_FILE)
    print(f"\n✓  報告已輸出：{OUT_FILE}")
    print(f"   檔案大小：{os.path.getsize(OUT_FILE)/1024:.0f} KB")


if __name__ == '__main__':
    print("=" * 60)
    print("亞朵 (ATAT) Q1 2026 業績更新 — DOCX 報告生成")
    print("=" * 60)
    build()
    print("=" * 60)
