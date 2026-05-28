"""
PDD Holdings 2026 年第一季度盈利更新報告（繁體中文版）
輸出：PDD_Q1_2026_盈利更新報告.docx
機構級格式：約 11 頁，使用繁體中文（標楷體/微軟正黑體），嵌入 10 張中文圖表。
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHARTS = os.path.join(os.path.dirname(__file__), 'charts_zh')
OUT_FILE = os.path.join(os.path.dirname(__file__), 'PDD_Q1_2026_盈利更新報告.docx')

# 中文字體設定
ZH_FONT = '微軟正黑體'  # Microsoft JhengHei，Windows/macOS Office 都有
ZH_FONT_EAST = 'PMingLiU'  # 細明體 fallback

# 顏色
NAVY   = RGBColor(0x1B, 0x4F, 0x8A)
ORANGE = RGBColor(0xE8, 0x72, 0x2A)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREEN  = RGBColor(0x2E, 0x8B, 0x57)
GRAY   = RGBColor(0x7F, 0x8C, 0x8D)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)


def set_cell_bg(cell, rgb_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)


def set_run_chinese_font(run, font_name=ZH_FONT):
    """為 run 設定中文字體（同時設定 East Asian font）"""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def add_heading(doc, text, level=1, color=NAVY, size=None):
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    set_run_chinese_font(run)
    run.font.size = Pt(size or sizes.get(level, 11))
    run.font.color.rgb = color
    pPr = p._p.get_or_add_pPr()
    if level == 1:
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '1B4F8A')
        pBdr.append(bottom)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, indent=False, space_after=4, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    set_run_chinese_font(run)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' ')
        r1.bold = True
        set_run_chinese_font(r1)
        r1.font.size = Pt(10)
    run = p.add_run(text)
    set_run_chinese_font(run)
    run.font.size = Pt(10)
    return p


def add_chart(doc, filename, width=6.0, caption=None):
    path = os.path.join(CHARTS, filename)
    if not os.path.exists(path):
        print(f"  警告：找不到圖表 {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_chinese_font(r)
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        r.italic = True
        cap.paragraph_format.space_after = Pt(6)


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.0):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_page_break(doc):
    doc.add_page_break()


def divider(doc, color='1B4F8A'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '1B4F8A')
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(color_el)
    rPr.append(u_el)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ZH_FONT)
    rFonts.set(qn('w:hAnsi'), ZH_FONT)
    rFonts.set(qn('w:eastAsia'), ZH_FONT)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)


# ═══════════════════════════════════════════════════════════════════
# 報告組裝
# ═══════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_page_margins(doc)

    # 設定預設樣式為中文字體
    normal = doc.styles['Normal']
    normal.font.name = ZH_FONT
    normal.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════
    # 第一頁：封面與摘要
    # ══════════════════════════════════════════════════════════════

    # 標題藍色橫條
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '1B4F8A')

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('盈利更新報告')
    set_run_chinese_font(r)
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('PDD Holdings Inc.（NASDAQ：PDD）')
    set_run_chinese_font(r2)
    r2.font.size = Pt(18)
    r2.bold = True
    r2.font.color.rgb = WHITE

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('2026 年第一季度業績 ｜ 2026 年 5 月 27 日')
    set_run_chinese_font(r3)
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1)
    r3.italic = True

    for pp in cell.paragraphs:
        pp.paragraph_format.space_after = Pt(4)
    cell.paragraphs[0].paragraph_format.space_before = Pt(8)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 評級卡（4 格橫排）
    rt = doc.add_table(rows=1, cols=4)
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ['投資評級', '新目標股價', '原目標股價', '現行股價']
    values = ['買入', '美元 125', '美元 165', '約 86 美元']
    colors_bg = ['1B4F8A', 'E8722A', '7F8C8D', '2C3E50']
    for cell, lbl, val, bg in zip(rt.row_cells(0), labels, values, colors_bg):
        set_cell_bg(cell, bg)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p1.add_run(lbl)
        set_run_chinese_font(rl)
        rl.font.size = Pt(9)
        rl.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = p2.add_run(val)
        set_run_chinese_font(rv)
        rv.font.size = Pt(14)
        rv.bold = True
        rv.font.color.rgb = WHITE
        p1.paragraph_format.space_before = Pt(6)
        p2.paragraph_format.space_after = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 核心結論
    add_heading(doc, '核心結論 — 2026 年第一季度業績要點', level=1)

    takeaways = [
        ('營收遜預期 3.4%',
         '本季營業收入為人民幣 1,062 億元（154 億美元），年增 11%，較 Bloomberg 共識預估的 159.4 億美元低約 5.4 億美元。營收減速主要反映 Temu 由純市場平台向第一方品牌模式轉型，以及主要海外市場面臨的宏觀逆風。'),
        ('每股盈利大幅遜預期 35%',
         '非 GAAP 攤薄每股盈利為人民幣 9.51 元（1.38 美元），低於共識預估的 2.13 美元，差距達 0.75 美元或 35.2%。管理層刻意壓制短期盈利，以資助為期三年、總額人民幣 1,000 億元（138 億美元）的供應鏈與第一方品牌投資計畫。'),
        ('淨利潤年減 15%',
         'GAAP 淨利潤降至人民幣 125 億元（年減 15%），儘管營業利潤年增 22% 至人民幣 196 億元。營業利潤與淨利潤走勢的背離，反映營業利潤線以下的投資支出大幅增加，以及實際稅率上升。'),
        ('交易服務收入加速增長',
         '交易服務收入年增 20% 至人民幣 563 億元，佔總營收比重達 53%（年增 400 個基點）。此走勢顯示平台貨幣化能力持續深化，即使線上行銷服務收入增速僅 2.5%。'),
        ('戰略轉型：第一方品牌計畫',
         '管理層於 2026 年 3 月成立專責第一方品牌業務的子公司，承諾未來三年投入人民幣 1,000 億元。此舉複製亞馬遜「平台 → 自有品牌」的長期戰略，以強化供應鏈控制力與長期利潤率。'),
        ('維持買入評級｜目標股價由 $165 下修至 $125',
         '我們將 12 個月目標股價由 165 美元下修至 125 美元（下調 24%），反映短期盈利受壓而調降的估算。但維持買入評級，因為現價約 86 美元至目標股價隱含 45% 潛在升幅，且供應鏈投資論述的長期邏輯依然成立。'),
    ]

    for bold_text, body_text in takeaways:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        r1 = p.add_run('■  ')
        r1.font.color.rgb = NAVY
        r1.font.size = Pt(10)
        r2 = p.add_run(bold_text + '：')
        r2.bold = True
        set_run_chinese_font(r2)
        r2.font.size = Pt(10)
        r2.font.color.rgb = NAVY
        r3 = p.add_run(body_text)
        set_run_chinese_font(r3)
        r3.font.size = Pt(10)

    divider(doc)

    # 季度業績摘要表
    add_heading(doc, '季度業績摘要表', level=2)

    summary_data = [
        ['指標', '2025 Q1', '2026 Q1', '年增率', '共識預估*', '超/遜預期'],
        ['營業收入（人民幣十億元）', '95.7', '106.2', '+11.0%', '109.3', '−3.4%'],
        ['— 線上行銷服務', '48.7', '49.9', '+2.5%', '—', '—'],
        ['— 交易服務', '47.0', '56.3', '+19.8%', '—', '—'],
        ['毛利（人民幣十億元）', '54.8', '59.3', '+8.2%', '—', '—'],
        ['毛利率', '57.3%', '55.8%', '−150bps', '—', '—'],
        ['營業利潤（人民幣十億元）', '16.1', '19.6', '+21.7%', '—', '—'],
        ['營業利潤率', '16.8%', '18.5%', '+170bps', '—', '—'],
        ['淨利潤（人民幣十億元）', '14.7', '12.5', '−15.0%', '—', '—'],
        ['淨利率', '15.4%', '11.8%', '−360bps', '—', '—'],
        ['非GAAP淨利潤（人民幣十億元）', '16.9', '14.1', '−16.6%', '—', '—'],
        ['非GAAP攤薄EPS（人民幣元/ADS）', '11.41', '9.51', '−16.6%', '15.47', '−38.5%'],
        ['非GAAP攤薄EPS（美元/ADS）', '$1.57', '$1.38', '−12.1%', '$2.13', '−35.2%'],
        ['營業現金流（人民幣十億元）', '15.5', '16.4', '+5.8%', '—', '—'],
        ['現金及短期投資（人民幣十億元）', '—', '436.1', '—', '—', '—'],
    ]

    tbl2 = doc.add_table(rows=len(summary_data), cols=6)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(summary_data):
        row = tbl2.rows[i]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_run_chinese_font(r)
            r.font.size = Pt(8.5)
            if i == 0:
                r.bold = True
                r.font.color.rgb = WHITE
                set_cell_bg(cell, '1B4F8A')
            elif i % 2 == 0:
                set_cell_bg(cell, 'EBF5FB')
            if j == 0 and i > 0:
                r.bold = True
            if '−38' in val or '−35' in val or '−3.4' in val:
                r.font.color.rgb = RED
                r.bold = True

    p_note = doc.add_paragraph()
    r_note = p_note.add_run('* Bloomberg 共識預估為 2026 年 5 月 26 日（業績發布前一日）數據')
    set_run_chinese_font(r_note)
    r_note.font.size = Pt(7.5)
    r_note.italic = True
    r_note.font.color.rgb = GRAY

    # ══════════════════════════════════════════════════════════════
    # 第二章：業績詳細分析
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, '一、業績詳細分析', level=1)

    add_heading(doc, '1.1 營業收入：年增 11%，低於市場預期', level=2)
    add_body(doc, 'PDD Holdings 公布 2026 年第一季度營業收入為人民幣 1,062 億元（154 億美元），年增 11.0%。實際結果較 Bloomberg 共識預估的 159.4 億美元低 5.4 億美元（3.4%），這是該公司連續第二個季度營收遜於市場預期。本次遜預期主要反映 Temu 國際業務模式的結構性轉變，而非景氣循環性的需求疲弱。')
    add_body(doc, '交易服務收入達人民幣 563 億元（年增 19.8%），為本季營收的主要動能，反映平台對物流履約、支付處理及供應鏈服務的貨幣化能力深化。該分部目前佔總營收 53.0%，較 2025 年第一季度的 49.1% 提升 400 個基點。相對而言，線上行銷服務收入僅年增 2.5% 至人民幣 499 億元，反映：（1）Temu 由高成本外部付費行銷模式轉向以自有供應鏈效率作為主要增長動能；（2）拼多多國內平台因抖音（字節跳動）與快手競爭壓力加劇，廣告需求趨緩。')

    add_chart(doc, 'chart_01_季度營收走勢.png', width=5.8,
              caption='圖一：PDD Holdings 季度營業收入走勢（人民幣十億元）— 2024 Q1 至 2026 Q1\n資料來源：PDD Holdings 公告；分析師估算 2025 Q2-Q4 數據')
    add_chart(doc, 'chart_02_分部營收拆解.png', width=5.8,
              caption='圖二：分部營收拆解 — 交易服務佔比攀升至 53%\n資料來源：PDD Holdings 2026 Q1 業績發布（2026年5月27日）；Form 6-K')

    add_heading(doc, '1.2 盈利能力：營運槓桿改善；淨利潤受投資支出影響', level=2)
    add_body(doc, '本季每股盈利遜預期的表面結果，掩蓋了一項重要事實：本季營業利潤水平實質改善。營業利潤達人民幣 196 億元（年增 21.7%），營業利潤率提升 170 個基點至 18.5%。改善動能來自物流與科技基礎設施可變成本效率提升，但部分被銷售成本上升所抵消（年增 14.7% 至人民幣 469 億元），原因包括履約成本、支付處理費用及頻寬支出增加。')
    add_body(doc, '然而，GAAP 淨利潤年減 15.0% 至人民幣 125 億元，非 GAAP 淨利潤年減 16.6% 至人民幣 141 億元。營業利潤（年增 22%）與淨利潤（年減 15%）走勢的鮮明背離，反映以下三項位於營業利潤線以下的因素：（1）部分中國境內主體的稅收優惠到期，所得稅準備金增加；（2）對新成立的第一方品牌主體進行的金融投資於營業利潤線以下入帳；（3）股權激勵攤銷費用上升。')

    add_chart(doc, 'chart_04_利潤率走勢.png', width=5.8,
              caption='圖三：季度利潤率走勢 — 毛利率穩定；淨利率年減 360 個基點\n資料來源：PDD Holdings 2026 Q1 業績發布（2026年5月27日）；Form 6-K')

    # ══════════════════════════════════════════════════════════════
    # 第三章：關鍵指標與管理層指引
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, '二、關鍵指標與管理層指引', level=1)

    add_heading(doc, '2.1 每股盈利：刻意遜預期，並非業務惡化', level=2)
    add_body(doc, '本季非 GAAP 攤薄每股盈利為人民幣 9.51 元（1.38 美元），年減 16.6%，較 Bloomberg 共識預估的 2.13 美元低 35.2%，創下 PDD 上市以來最大幅度的季度每股盈利遜預期紀錄。我們需要向客戶強調：這是一項戰略選擇，並非業務根基惡化的信號。')
    add_body(doc, '管理層於業績會議的措辭明確無誤：「我們不追求短期財務表現的最優化，而是優先考慮平台生態系統的健康發展。」公司透過下列管道進行投資：（1）承諾未來三年投入人民幣 1,000 億元發展第一方品牌（每年約人民幣 333 億元，相當於目前年度營收的約 31%）；（2）對偏遠鄉鎮提供免費物流服務（試點地區覆蓋率已逾 70%）；（3）本季新增逾 20 項食品安全合規措施；（4）直接補貼農業供應鏈。')

    add_chart(doc, 'chart_03_每股盈利走勢.png', width=5.8,
              caption='圖四：季度每股盈利走勢（GAAP 與非 GAAP，每股 ADS）— 2024 Q1 至 2026 Q1\n資料來源：PDD Holdings 2026 Q1 業績發布；Bloomberg 共識預估')
    add_chart(doc, 'chart_05_超預期遜預期摘要.png', width=5.5,
              caption='圖五：2026 Q1 實際業績 vs Bloomberg 共識預估\n資料來源：Bloomberg 共識預估（2026 年 5 月 26 日）；PDD Holdings 2026 Q1 業績發布')

    add_heading(doc, '2.2 收入結構轉變：結構性而非景氣循環性', level=2)
    add_body(doc, '交易服務收入加速增長（年增 20%）相對於線上行銷收入趨緩（年增 2.5%），是管理層刻意的戰略訊號。交易服務收入包含 Temu 履約費用、國內物流網路費用及商家支付處理費用，會隨 PDD 自有基礎設施擴張而增長。此轉變降低了 PDD 對付費廣告投資回報週期的依賴，並建立更可持續的、以資產為基礎的收入流。我們認為交易服務佔比提升從長期看有利於收入質量，雖然在基礎設施建設期會壓抑利潤率。')

    add_chart(doc, 'chart_06_營收年增率.png', width=5.8,
              caption='圖六：季度營收年增率走勢 — 先減速後回升\n資料來源：PDD Holdings 2026 Q1 業績發布；Bloomberg 共識預估')

    add_heading(doc, '2.3 成本結構與營運槓桿', level=2)
    add_body(doc, '銷售成本於 2026 年第一季度年增 14.7% 至人民幣 469 億元，略高於營收增速 11.0%，導致毛利率壓縮 150 個基點。但營業費用（不含銷售成本）控制良好：銷售與行銷費用僅年增 9.3% 至人民幣 224 億元（較 2023-2024 年 Temu 全球擴張期的行銷支出明顯放緩），研發費用年減 7.7% 至人民幣 132 億元。綜合結果是營業利潤率提升 170 個基點，顯示平台底層效率持續改善。')

    add_chart(doc, 'chart_07_營運開支對比.png', width=5.8,
              caption='圖七：損益表瀑布圖 — 2025 Q1 vs 2026 Q1（人民幣十億元）\n資料來源：PDD Holdings 2026 Q1 業績發布（2026年5月27日）；Form 6-K')

    add_heading(doc, '2.4 資產負債表與現金流', level=2)
    add_body(doc, 'PDD 維持堡壘級資產負債表，現金及短期投資達人民幣 4,361 億元（632 億美元），相當於每股 ADS 約 41.70 美元，約佔業績後股價的 49%。本季營業現金流為人民幣 164 億元（年增 5.8%），現金生成能力依然強勁。公司獨立財務報表上無長期負債。如此龐大的現金儲備為下列三項用途提供巨大彈性：人民幣 1,000 億元投資計畫的執行、股票回購、以及任何與 Temu 合規調查相關的監管處理成本。')

    add_heading(doc, '2.5 業績展望：公司政策不提供量化指引', level=2)
    add_body(doc, '依公司長期政策，PDD 不提供量化的營收或盈利預測指引。管理層於 2026 年第一季度業績會議的定性指引可歸納如下：')
    guidance_items = [
        ('投資週期延續：', '人民幣 1,000 億元第一方品牌承諾為多年期計畫。客戶不應預期短期（2026 Q2-Q3）每股盈利反彈。'),
        ('Temu 模式轉型：', '「我們將深化全球市場的第一方品牌模式發展。」這證實 Temu 由純市場平台向自有庫存/品牌模式轉型，初期聚焦核心 SKU 類別。'),
        ('農業投資：', '持續優先發展農業直連供應鏈、鄉村物流及食品安全合規。本季新增逾 20 項食品安全計畫。'),
        ('監管合規：', '管理層承認監管審查持續，但未量化合規成本。表態將「積極配合」所有相關監管機構。'),
    ]
    for prefix, body in guidance_items:
        add_bullet(doc, body, bold_prefix=prefix)

    # ══════════════════════════════════════════════════════════════
    # 第四章：更新後投資論述
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, '三、更新後投資論述', level=1)

    add_heading(doc, '3.1 投資論述基礎依然完整，但短期盈利受壓', level=2)
    add_body(doc, '我們的首次覆蓋投資論述建立在三大支柱之上：（1）PDD 在中國的 C2M 供應鏈領導地位；（2）Temu 全球擴張軌跡；（3）公司將上述優勢轉化為持續超越共識預估盈利能力的執行力。在 2026 年第一季度業績發布後，我們重新審視每一項支柱。')

    thesis_points = [
        ('支柱一 — 中國供應鏈領導地位：完整保留。',
         '國內拼多多平台持續產生強勁的交易服務收入增長（合併年增 20%），證實供應鏈基礎設施貨幣化能力持續深化。農業與工業 C2M 計畫持續擴張。我們未觀察到核心國內業務有任何結構性惡化。'),
        ('支柱二 — Temu 全球擴張：轉型中。',
         '業務敘事由「超高速增長市場平台」轉為「第一方品牌與供應鏈能力建設者」。此戰略合理 — Temu 面臨美國 de minimis 免稅額度的監管壓力及關稅風險 — 但將導致數個季度的營收增速壓縮及資本支出增加，而後相關效益才能顯現。我們將 Temu 至 2028 年的營收複合增長率預估由 25% 下修至 18%。'),
        ('支柱三 — 超額盈利能力路徑：延後，並非放棄。',
         '管理層正進行刻意而大規模的投資，將延後盈利能力的復甦。我們現預期非 GAAP 營業利潤率將於 2026 年上半年觸底（約 17-19%），隨後於 2027-2028 年隨第一方品牌產生營運槓桿而回升至 22-25%。每股盈利復甦將自 2027 年起，隨投資正常化而展開。'),
    ]

    for bold_text, body_text in thesis_points:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.3
        r1 = p.add_run(bold_text + ' ')
        r1.bold = True
        set_run_chinese_font(r1)
        r1.font.size = Pt(10)
        r1.font.color.rgb = NAVY
        r2 = p.add_run(body_text)
        set_run_chinese_font(r2)
        r2.font.size = Pt(10)

    add_heading(doc, '3.2 主要風險（更新版）', level=2)

    risks = [
        ('監管升級風險（高）：', 'Temu 面臨美國（CPSC 產品安全）、歐盟（數位服務法合規）及中國（CNIPA）多方持續調查。任何重大罰款或平台限制均將構成盈利重大逆風。'),
        ('投資週期超出預期（中）：', '若人民幣 1,000 億元承諾前置加速或大幅超出原規劃，短期每股盈利壓制可能延續至 2027 年下半年。'),
        ('線上行銷收入停滯（中）：', '行銷服務 2.5% 年增率暗示拼多多國內廣告收入可能趨近成熟。若停滯延續，將限制長期營收增長空間。'),
        ('關稅與貿易政策（中）：', '美國 301 條款關稅及低值包裹 de minimis 免稅額度到期，對 Temu 原市場模式構成結構性逆風。'),
        ('中國宏觀疲軟（中低）：', '中國消費者信心依然脆弱。國內消費較預期更深的放緩將影響拼多多 GMV 增長。'),
    ]

    for prefix, body in risks:
        add_bullet(doc, body, bold_prefix=prefix)

    # ══════════════════════════════════════════════════════════════
    # 第五章：估值與估算更新
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, '四、估值與估算更新', level=1)

    add_heading(doc, '4.1 估算修訂', level=2)
    add_body(doc, '我們大幅下修營收與盈利估算，反映：（甲）2026 年第一季度確認的每股盈利執行水平；（乙）管理層對近期持續投資的定性指引；（丙）我們對 Temu 至 2028 年營收增速假設的調降。下表彙總修訂後估算。')

    est_data = [
        ['', 'FY2025E', 'FY2026E（原）', 'FY2026E（新）', 'Δ%', 'FY2027E（原）', 'FY2027E（新）', 'Δ%'],
        ['營收（人民幣十億元）', '421', '450', '432', '−4.0%', '510', '498', '−2.4%'],
        ['年增率', '+10.5%', '+7.0%', '+2.6%', '—', '+13.3%', '+15.3%', '—'],
        ['毛利率', '56.1%', '56.5%', '55.5%', '−100bps', '57.0%', '56.5%', '−50bps'],
        ['非GAAP營業利潤率', '18.3%', '21.5%', '18.8%', '−270bps', '23.0%', '21.5%', '−150bps'],
        ['非GAAP淨利潤（人民幣十億元）', '48.2', '67.5', '45.0', '−33.3%', '84.0', '59.0', '−29.8%'],
        ['非GAAP EPS（人民幣元/ADS）', '46.8', '65.5', '43.7', '−33.3%', '81.5', '57.3', '−29.7%'],
        ['非GAAP EPS（美元/ADS）', '$6.45', '$9.02', '$6.02', '−33.3%', '$11.23', '$7.89', '−29.7%'],
    ]

    tbl3 = doc.add_table(rows=len(est_data), cols=8)
    tbl3.style = 'Table Grid'
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(est_data):
        for j, val in enumerate(row_data):
            cell = tbl3.cell(i, j)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_run_chinese_font(r)
            r.font.size = Pt(8)
            if i == 0:
                r.bold = True
                r.font.color.rgb = WHITE
                set_cell_bg(cell, '1B4F8A')
            elif i % 2 == 0:
                set_cell_bg(cell, 'EBF5FB')
            if j == 0 and i > 0:
                r.bold = True
            if '−33' in val or '−29' in val or '−27' in val or '−4.0' in val:
                r.font.color.rgb = RED

    add_chart(doc, 'chart_08_估算修訂.png', width=5.8,
              caption='圖八：營收與每股盈利估算修訂對比 — 原估算 vs 新估算\n資料來源：分析師估算；PDD Holdings 2026 Q1 業績發布（2026 年 5 月 27 日）')

    add_heading(doc, '4.2 目標股價修訂：由 $165 下修至 $125', level=2)
    add_body(doc, '我們將 12 個月目標股價由 165 美元下修至 125 美元，反映調降後的盈利估算及短期盈利不確定性下適度壓縮的目標倍數。本目標股價採用混合方法論推導：')

    pt_data = [
        ['估值方法', '估算基礎', '倍數/折現率', '隱含股價', '權重', '貢獻'],
        ['本益比（FY2027E 非GAAP）', '$7.89/ADS', '16.0倍', '$126', '40%', '$50'],
        ['EV/EBITDA（FY2027E）', '$10.2B EBITDA', '12.0倍', '$118', '30%', '$35'],
        ['現金流折現（10年，WACC 10.5%）', '8.5% 終值增長', 'g = 3%', '$145', '20%', '$29'],
        ['分部加總（PDD+Temu）', '分部混合', 'PDD 12倍 + Temu 8倍', '$138', '10%', '$14'],
        ['混合目標股價', '', '', '', '100%', '$128'],
        ['四捨五入後目標股價', '', '', '', '', '$125'],
    ]

    tbl4 = doc.add_table(rows=len(pt_data), cols=6)
    tbl4.style = 'Table Grid'
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(pt_data):
        for j, val in enumerate(row_data):
            cell = tbl4.cell(i, j)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_run_chinese_font(r)
            r.font.size = Pt(8.5)
            if i == 0:
                r.bold = True
                r.font.color.rgb = WHITE
                set_cell_bg(cell, '1B4F8A')
            elif i >= len(pt_data) - 2:
                set_cell_bg(cell, 'FFF3CD')
                r.bold = True
                r.font.color.rgb = NAVY
            elif i % 2 == 0:
                set_cell_bg(cell, 'EBF5FB')
            if j == 0 and i > 0:
                r.bold = True

    add_body(doc, '在我們修訂後的目標股價 125 美元水平，PDD 將以 FY2027E 非 GAAP 每股盈利的 15.8 倍交易，相對亞馬遜（本益比 35 倍）有顯著折讓，與 Sea（本益比 22 倍）相當。考慮短期投資週期與監管不確定性，此折讓水平合理。在現價約 86 美元水平，至我們基本情境目標股價的隱含潛在升幅為 +45%，足以支持買入評級。', space_after=6)

    add_chart(doc, 'chart_09_同業估值比較.png', width=5.8,
              caption='圖九：業績後本益比與 EV/EBITDA 同業比較\n資料來源：Bloomberg；分析師估算（2026 年 5 月）；PDD Holdings 2026 Q1 業績發布')
    add_chart(doc, 'chart_10_目標股價情境分析.png', width=5.5,
              caption='圖十：目標股價情境分析 — 樂觀/基本/悲觀\n資料來源：分析師估算；PDD Holdings 2026 Q1 業績發布（2026 年 5 月 27 日）')

    # ══════════════════════════════════════════════════════════════
    # 第六章：來源與披露
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, '資料來源與參考文獻', level=1)

    add_heading(doc, '主要資料來源 — 2026 年第一季度業績相關文件', level=2)

    sources = [
        ('PDD Holdings 2026 年第一季度業績發布（2026 年 5 月 27 日）',
         'https://investor.pddholdings.com/',
         '公司公布未經審計的 2026 年第一季度財務業績的新聞稿。本報告所引用的全部營收、利潤及每股盈利數據之主要資料來源。'),
        ('PDD Holdings Form 6-K 文件（2026 年 5 月 27 日呈報）',
         'https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001737806&type=6-K&dateb=&owner=include&count=40',
         '向美國證券交易委員會（SEC）EDGAR 系統呈報的 2026 年第一季度未經審計財務業績文件。文件編號：0001104659-26-034813。'),
        ('PDD Holdings 2026 年第一季度業績說明會逐字稿（2026 年 5 月 27 日）',
         'https://www.benzinga.com/insights/news/26/05/52804676/pdd-holdings-q1-2026-earnings-call-complete-transcript',
         '管理層業績說明會發言及分析師問答環節完整逐字稿。本報告所引用的全部管理層發言之資料來源。'),
        ('Bloomberg 共識預估（2026 年 5 月 26 日）',
         'https://www.bloomberg.com/quote/PDD:US',
         '營收共識預估 159.4 億美元；非 GAAP 每股盈利共識預估 2.13 美元/ADS。本報告全部超預期/遜預期計算之基準。'),
    ]

    for i, (title, url, desc) in enumerate(sources):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        r_num = p.add_run(f'[{i+1}]  ')
        r_num.bold = True
        set_run_chinese_font(r_num)
        r_num.font.size = Pt(9)
        r_num.font.color.rgb = NAVY
        add_hyperlink(p, title, url)
        r_desc = p.add_run(f' — {desc}')
        set_run_chinese_font(r_desc)
        r_desc.font.size = Pt(9)
        r_desc.font.color.rgb = GRAY

    add_heading(doc, '輔助參考資料', level=2)

    supp = [
        ('StockTitan — PDD Q1 2026 業績摘要',
         'https://www.stocktitan.net/news/PDD/pdd-holdings-announces-first-quarter-2026-unaudited-financial-2iwyuhtjni8p.html'),
        ('PDD Holdings 投資人關係官網',
         'https://investor.pddholdings.com/'),
        ('PDD Holdings SEC EDGAR 文件',
         'https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001737806&type=&dateb=&owner=include&count=40'),
    ]

    for title, url in supp:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(2)
        r_bullet = p.add_run('•  ')
        r_bullet.font.size = Pt(9)
        r_bullet.font.color.rgb = NAVY
        add_hyperlink(p, title, url)

    divider(doc)

    add_heading(doc, '分析師認證與披露聲明', level=2)
    disc_text = (
        '本報告所表達之觀點，係負責本報告之分析師個人對所涉及的有價證券及/或發行機構的真實看法。'
        '分析師薪酬之任何部分，無論直接或間接，均與本報告所載之具體建議或觀點無關。本報告僅供資訊參考之用，'
        '不構成任何有價證券買賣之要約或邀請。過往業績不代表未來表現。投資人應自行進行盡職調查。'
        '本盈利更新報告於 2026 年 5 月 27 日業績發布後 24 小時內完成，符合機構級研究標準。'
    )
    p_disc = doc.add_paragraph()
    p_disc.paragraph_format.left_indent = Inches(0.1)
    p_disc.paragraph_format.line_spacing = 1.4
    r_disc = p_disc.add_run(disc_text)
    set_run_chinese_font(r_disc)
    r_disc.font.size = Pt(8)
    r_disc.font.color.rgb = GRAY
    r_disc.italic = True

    # 頁尾
    for section in doc.sections:
        footer = section.footer
        ft_para = footer.paragraphs[0]
        ft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ft_run = ft_para.add_run(
            'PDD Holdings Inc.（NASDAQ：PDD）— 2026 年第一季度盈利更新報告  ｜  '
            '2026 年 5 月 27 日  ｜  投資評級：買入  ｜  目標股價：美元 125  ｜  '
            '機密文件 — 僅供專業投資者使用'
        )
        set_run_chinese_font(ft_run)
        ft_run.font.size = Pt(7.5)
        ft_run.font.color.rgb = GRAY

    doc.save(OUT_FILE)
    size_kb = os.path.getsize(OUT_FILE) / 1024
    print(f"\n  ✓  {OUT_FILE}  ({size_kb:.0f} KB)")


if __name__ == '__main__':
    print("開始組裝 PDD Holdings 2026 Q1 盈利更新報告（繁體中文版）...")
    build()
    print("完成。")
