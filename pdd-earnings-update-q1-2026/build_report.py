"""
PDD Holdings Q1 2026 — Earnings Update Report Builder
Generates: PDD_Q1_2026_Earnings_Update.docx
Institutional format: 10 pages, Times New Roman, 10 charts embedded.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

CHARTS = os.path.join(os.path.dirname(__file__), 'charts')
OUT_FILE = os.path.join(os.path.dirname(__file__), 'PDD_Q1_2026_Earnings_Update.docx')

# ─── Colour palette ───────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x4F, 0x8A)
ORANGE = RGBColor(0xE8, 0x72, 0x2A)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREEN  = RGBColor(0x2E, 0x8B, 0x57)
GRAY   = RGBColor(0x7F, 0x8C, 0x8D)
LGRAY  = RGBColor(0xF2, 0xF3, 0xF4)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
GOLD   = RGBColor(0xD4, 0xAC, 0x0D)

# ─── Helpers ──────────────────────────────────────────────────────

def set_cell_bg(cell, rgb_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)


def set_cell_borders(cell, border_color='1B4F8A', size=6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:color'), border_color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para_run(para, text, bold=False, italic=False, size=10, color=None, font='Times New Roman'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1, color=NAVY, size=None):
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size or sizes.get(level, 11))
    run.font.color.rgb = color
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    if level == 1:
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '1B4F8A')
        pBdr.append(bottom)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, indent=False, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' ')
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(10)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    return p


def add_chart(doc, filename, width=6.0, caption=None):
    path = os.path.join(CHARTS, filename)
    if not os.path.exists(path):
        print(f"  WARNING: chart not found: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        r.italic = True
        cap.paragraph_format.space_after = Pt(6)


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.0):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_page_break(doc):
    doc.add_page_break()


def divider(doc, color='1B4F8A'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    # Color and underline
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '1B4F8A')
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(color_el)
    rPr.append(u_el)
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)
    return hyperlink


# ═══════════════════════════════════════════════════════════════════
# REPORT BUILDER
# ═══════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_page_margins(doc)

    # ── Default Normal style ──────────────────────────────────────
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════
    # PAGE 1: COVER / SUMMARY
    # ══════════════════════════════════════════════════════════════

    # Header bar (blue table, 1 row, 1 col)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '1B4F8A')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('EARNINGS UPDATE')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = WHITE
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('PDD Holdings Inc. (NASDAQ: PDD)')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(18)
    r2.bold = True
    r2.font.color.rgb = WHITE
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('Q1 2026 Results | May 27, 2026')
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1)
    r3.italic = True
    cell.paragraphs[0].paragraph_format.space_before = Pt(8)
    for pp in cell.paragraphs:
        pp.paragraph_format.space_after = Pt(4)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Rating / Price Target box ─────────────────────────────────
    rt = doc.add_table(rows=1, cols=4)
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ['RATING', 'PRICE TARGET (NEW)', 'PRIOR PRICE TARGET', 'CURRENT PRICE']
    values = ['BUY', 'US$125', 'US$165', '~US$86']
    colors_bg = ['1B4F8A', 'E8722A', '7F8C8D', '2C3E50']
    for i, (cell, lbl, val, bg) in enumerate(zip(rt.row_cells(0), labels, values, colors_bg)):
        set_cell_bg(cell, bg)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p1.add_run(lbl)
        rl.font.name = 'Times New Roman'
        rl.font.size = Pt(8)
        rl.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = p2.add_run(val)
        rv.font.name = 'Times New Roman'
        rv.font.size = Pt(14)
        rv.bold = True
        rv.font.color.rgb = WHITE
        p1.paragraph_format.space_before = Pt(6)
        p2.paragraph_format.space_after = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Key takeaways box ────────────────────────────────────────
    add_heading(doc, 'KEY TAKEAWAYS — Q1 2026 RESULTS', level=1)

    takeaways = [
        ('REVENUE MISS (−3.4%)', 'Revenue of RMB106.2B (US$15.4B) grew 11% YoY, missing Bloomberg consensus of US$15.94B by ~$540M. Revenue deceleration reflects Temu model transition from marketplace to First Party Brand and macro headwinds in key markets.'),
        ('SIGNIFICANT EPS MISS (−35%)', 'Non-GAAP diluted EPS of RMB9.51 (US$1.38) missed consensus of US$2.13 by $0.75 or 35.2%. Management is deliberately suppressing near-term EPS to fund a ¥100B (US$13.8B) three-year supply chain and First Party Brand initiative.'),
        ('NET INCOME −15% YoY', 'GAAP net income fell to RMB12.5B (-15% YoY) despite operating profit growing +22% YoY to RMB19.6B. The divergence reflects elevated investment expenditure and higher effective tax rates below the operating line.'),
        ('TRANSACTION SERVICES ACCELERATION', 'Transaction services revenue grew +20% YoY to RMB56.3B, reaching 53% of total revenue (+400bps YoY mix shift). This signals platform monetisation deepening even as marketing services growth stalls at +2.5% YoY.'),
        ('STRATEGIC PIVOT: FIRST PARTY BRAND', 'Management launched a dedicated First Party Brand entity in March 2026, committing RMB100B over three years. This replicates Amazon\'s playbook of platform → owned brand for supply chain control and margin enhancement over the long term.'),
        ('RATING MAINTAINED: BUY | PT CUT $165 → $125', 'We lower our price target from $165 to $125 (−24%) reflecting estimate cuts on near-term EPS suppression, but maintain BUY given 45% upside to PT at current ~$86 price and the long-term validity of the supply chain investment thesis.'),
    ]

    for bold_text, body_text in takeaways:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run('■  ')
        r1.font.color.rgb = NAVY
        r1.font.size = Pt(10)
        r2 = p.add_run(bold_text + ': ')
        r2.bold = True
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10)
        r2.font.color.rgb = NAVY
        r3 = p.add_run(body_text)
        r3.font.name = 'Times New Roman'
        r3.font.size = Pt(10)

    divider(doc)

    # ── Summary Financial Table ────────────────────────────────────
    add_heading(doc, 'QUARTERLY RESULTS SUMMARY', level=2)

    summary_data = [
        ['Metric', 'Q1 2025A', 'Q1 2026A', 'YoY Δ', 'Consensus Est.', 'Beat / Miss'],
        ['Total Revenue (RMB B)', '95.7', '106.2', '+11.0%', '109.3*', '−3.4%'],
        ['— Online Marketing Services', '48.7', '49.9', '+2.5%', '—', '—'],
        ['— Transaction Services', '47.0', '56.3', '+19.8%', '—', '—'],
        ['Gross Profit (RMB B)', '54.8', '59.3', '+8.2%', '—', '—'],
        ['Gross Margin', '57.3%', '55.8%', '−150bps', '—', '—'],
        ['Operating Profit (RMB B)', '16.1', '19.6', '+21.7%', '—', '—'],
        ['Operating Margin', '16.8%', '18.5%', '+170bps', '—', '—'],
        ['Net Income (RMB B)', '14.7', '12.5', '−15.0%', '—', '—'],
        ['Net Margin', '15.4%', '11.8%', '−360bps', '—', '—'],
        ['Non-GAAP Net Income (RMB B)', '16.9', '14.1', '−16.6%', '—', '—'],
        ['Non-GAAP Diluted EPS/ADS (RMB)', '11.41', '9.51', '−16.6%', '15.47*', '−38.5%'],
        ['Non-GAAP Diluted EPS/ADS (US$)', '$1.57', '$1.38', '−12.1%', '$2.13*', '−35.2%'],
        ['Operating Cash Flow (RMB B)', '15.5', '16.4', '+5.8%', '—', '—'],
        ['Cash & Investments (RMB B)', '—', '436.1', '—', '—', '—'],
    ]

    tbl2 = doc.add_table(rows=len(summary_data), cols=6)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Inches(2.0), Inches(0.9), Inches(0.9), Inches(0.7), Inches(1.1), Inches(0.9)]

    for i, row_data in enumerate(summary_data):
        row = tbl2.rows[i]
        for j, (cell, val) in enumerate(zip(row.cells, row_data)):
            try:
                row.cells[j].width = col_widths[j]
            except Exception:
                pass
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(8.5)
            if i == 0:
                r.bold = True
                r.font.color.rgb = WHITE
                set_cell_bg(cell, '1B4F8A')
            elif i % 2 == 0:
                set_cell_bg(cell, 'EBF5FB')
            if j == 0 and i > 0:
                r.bold = True
            if 'MISS' in val or '−38' in val or '−35' in val or '−3.4' in val:
                r.font.color.rgb = RED
                r.bold = True
            if 'BEAT' in val:
                r.font.color.rgb = GREEN
                r.bold = True

    p_note = doc.add_paragraph()
    r_note = p_note.add_run('* Bloomberg consensus as of May 26, 2026 (day prior to earnings release).')
    r_note.font.name = 'Times New Roman'
    r_note.font.size = Pt(7.5)
    r_note.italic = True
    r_note.font.color.rgb = GRAY

    # ══════════════════════════════════════════════════════════════
    # PAGE 2–3: DETAILED RESULTS ANALYSIS
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, '1. DETAILED RESULTS ANALYSIS', level=1)

    add_heading(doc, '1.1 Revenue: 11% YoY Growth — Below Expectations', level=2)
    add_body(doc, 'PDD Holdings reported Q1 2026 total revenues of RMB106.2 billion (US$15.4 billion), representing 11.0% year-over-year growth. The result came in $540 million (3.4%) below Bloomberg consensus of US$15.94 billion, marking the second consecutive quarter of consensus revenue misses. The miss reflects a deliberate structural shift in the Temu international business model rather than cyclical demand weakness.')
    add_body(doc, 'Transaction services revenues of RMB56.3 billion (+19.8% YoY) were the growth engine, driven by deepened monetisation of logistics fulfilment, payment processing, and supply chain services. This segment now represents 53.0% of total revenue, up 400bps from 49.1% in Q1 2025. In contrast, online marketing services revenues grew a modest 2.5% YoY to RMB49.9 billion, reflecting: (1) Temu\'s transition away from high-cost external paid marketing toward proprietary supply chain efficiency as the primary growth lever; and (2) softening advertising demand on the Pinduoduo domestic platform amid increased competitive pressure from Douyin (ByteDance) and Kuaishou.')

    add_chart(doc, 'chart_01_quarterly_revenue.png', width=5.8,
              caption='Figure 1: PDD Holdings Quarterly Revenue (RMB B) — Q1 2024 to Q1 2026\nSource: PDD Holdings filings; analyst estimates for Q2–Q4 2025')
    add_chart(doc, 'chart_02_revenue_by_segment.png', width=5.8,
              caption='Figure 2: Revenue by Segment — Transaction Services Reaching 53% Mix\nSource: PDD Holdings Q1 2026 Earnings Release (May 27, 2026); Form 6-K')

    add_heading(doc, '1.2 Profitability: Operating Leverage Intact; Net Income Impacted by Investment', level=2)
    add_body(doc, 'The headline EPS miss masks an important distinction: operating profitability improved materially year-on-year. Operating profit of RMB19.6 billion grew 21.7% YoY, with operating margin expanding 170bps to 18.5%. This improvement reflects improved variable cost efficiency in logistics and technology infrastructure, partially offset by higher cost of revenues (+14.7% YoY to RMB46.9B), driven by fulfilment costs, payment processing fees, and bandwidth expenses.')
    add_body(doc, 'However, GAAP net income fell 15.0% YoY to RMB12.5 billion, and non-GAAP net income declined 16.6% to RMB14.1 billion. The sharp divergence between operating profit (+22%) and net income (−15%) reflects three factors below the operating line: (1) higher tax provisions as the preferential tax treatment for certain PRC entities phases out; (2) increased financial investments in the new First Party Brand entity recognised below the line; and (3) elevated stock-based compensation amortisation.')

    add_chart(doc, 'chart_04_margin_trends.png', width=5.8,
              caption='Figure 3: Quarterly Margin Trends — Gross Margin Stable; Net Margin Compressed 360bps YoY\nSource: PDD Holdings Q1 2026 Earnings Release (May 27, 2026); Form 6-K')

    # ══════════════════════════════════════════════════════════════
    # PAGE 4–5: KEY METRICS & GUIDANCE
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, '2. KEY METRICS & MANAGEMENT GUIDANCE', level=1)

    add_heading(doc, '2.1 EPS: Intentional Miss, Not a Business Deterioration', level=2)
    add_body(doc, 'Non-GAAP diluted EPS of RMB9.51 (US$1.38) per ADS declined 16.6% YoY and missed Bloomberg consensus of US$2.13 by 35.2% — the largest quarterly EPS miss in PDD\'s history as a public company. We want to emphasise to clients that this is a strategic choice, not a sign of business deterioration.')
    add_body(doc, 'Management\'s language on the earnings call was unambiguous: "Instead of focusing on short-term financial performance, we prioritise the healthy development of the platform ecosystem." The company is funding this through: (1) RMB100 billion committed to First Party Brand over three years (~RMB33B per year, or ~31% of current annual revenues); (2) free logistics to rural villages (covering >70% of villages in pilot regions); (3) >20 food safety compliance initiatives launched in Q1; and (4) direct agricultural supply chain subsidies.')

    add_chart(doc, 'chart_03_eps_trend.png', width=5.8,
              caption='Figure 4: Quarterly EPS Trend (GAAP & Non-GAAP per ADS) — Q1 2024 to Q1 2026\nSource: PDD Holdings Q1 2026 Earnings Release; Bloomberg consensus')
    add_chart(doc, 'chart_05_beat_miss.png', width=5.5,
              caption='Figure 5: Q1 2026 Beat/Miss vs. Bloomberg Consensus\nSource: Bloomberg consensus (May 26, 2026); PDD Holdings Q1 2026 Earnings Release')

    add_heading(doc, '2.2 Revenue Mix Shift: Structural, Not Cyclical', level=2)
    add_body(doc, 'The acceleration of transaction services (+20% YoY) relative to online marketing (+2.5% YoY) is a deliberate strategic signal. Transaction services revenues — which encompass Temu fulfilment fees, domestic logistics network charges, and merchant payment processing — grow as PDD builds proprietary infrastructure. This reduces PDD\'s dependence on paid advertising ROI cycles and creates more durable, asset-backed revenue streams. We view the growing transaction services mix as a long-term positive for revenue quality, even though it is currently margin-dilutive during the build-out phase.')

    add_chart(doc, 'chart_06_revenue_growth.png', width=5.8,
              caption='Figure 6: Quarterly Revenue YoY Growth Rate — Deceleration Then Re-Acceleration\nSource: PDD Holdings Q1 2026 Earnings Release; Bloomberg consensus')

    add_heading(doc, '2.3 Cost Structure & Operating Leverage', level=2)
    add_body(doc, 'Cost of revenues rose 14.7% YoY to RMB46.9B in Q1 2026, slightly ahead of revenue growth of 11.0%, creating 150bps of gross margin compression. However, operating expenses (ex-COGS) were well-controlled, with sales & marketing up only 9.3% YoY to RMB22.4B (a deceleration from the Temu global expansion marketing spend of 2023–2024), and R&D declining 7.7% YoY to RMB13.2B. The net result was operating leverage of +170bps at the operating margin level, a positive indicator of underlying platform efficiency.')

    add_chart(doc, 'chart_07_opex_comparison.png', width=5.8,
              caption='Figure 7: P&L Cost Waterfall — Q1 2025 vs. Q1 2026 (RMB Billion)\nSource: PDD Holdings Q1 2026 Earnings Release (May 27, 2026); Form 6-K')

    add_heading(doc, '2.4 Balance Sheet & Cash Flow', level=2)
    add_body(doc, 'PDD maintains a fortress balance sheet with RMB436.1 billion (US$63.2 billion) in cash and short-term investments — equivalent to approximately $41.70 per ADS, or roughly 49% of the post-earnings share price. Operating cash flow of RMB16.4 billion (+5.8% YoY) demonstrated continued strong cash generation. The company has zero long-term debt on a standalone basis. This cash pile provides significant optionality for the planned RMB100B investment programme, share buybacks, and any regulatory resolution costs related to the ongoing Temu compliance investigations.')

    add_heading(doc, '2.5 Guidance: No Quantitative Guidance (Per Policy)', level=2)
    add_body(doc, 'PDD does not provide quantitative revenue or earnings guidance, consistent with its long-standing policy. Management qualitative guidance from the Q1 2026 earnings call can be summarised as follows:')
    guidance_items = [
        ('Investment cycle continuation:', 'The RMB100B First Party Brand commitment is a multi-year initiative. Clients should not expect EPS recovery in the near term (Q2–Q3 2026).'),
        ('Temu model transformation:', '"We will deepen the development of the First Party Brand model in the global market." This confirms Temu is moving from pure marketplace to an owned-inventory / brand model, initially targeting key SKU categories.'),
        ('Agricultural investment:', 'Continued prioritisation of direct agricultural supply chains, rural logistics, and food safety compliance. Over 20 food safety programmes launched in Q1 alone.'),
        ('Regulatory compliance:', 'Management acknowledged ongoing regulatory scrutiny but did not quantify compliance costs. They indicated the company is "actively cooperating" with all relevant authorities.'),
    ]
    for prefix, body in guidance_items:
        add_bullet(doc, body, bold_prefix=prefix)

    # ══════════════════════════════════════════════════════════════
    # PAGE 6–7: UPDATED INVESTMENT THESIS
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, '3. UPDATED INVESTMENT THESIS', level=1)

    add_heading(doc, '3.1 Thesis Intact Despite Near-Term EPS Pressure', level=2)
    add_body(doc, 'Our investment thesis at initiation rested on three pillars: (1) PDD\'s C2M supply chain dominance in China; (2) Temu\'s global expansion trajectory; and (3) the company\'s ability to leverage these advantages toward sustained, above-consensus profitability. Following Q1 2026 results, we revisit each pillar.')

    thesis_points = [
        ('Pillar 1 — China Supply Chain Dominance: INTACT.',
         'The domestic Pinduoduo platform continues to generate robust transaction services revenue growth (+20% YoY consolidated), confirming that monetisation of the supply chain infrastructure is deepening. Agricultural and industrial C2M programmes are expanding. We see no structural deterioration in the core domestic franchise.'),
        ('Pillar 2 — Temu Global Expansion: TRANSITIONING.',
         'The narrative has shifted from "hyper-growth marketplace" to "First Party Brand and supply chain capability builder." This is strategically sound — Temu faced regulatory pressure on the de minimis exemption in the US and tariff risks — but it will result in a multi-quarter period of compressed revenue growth and elevated capital expenditure before benefits materialise. We reduce our Temu revenue CAGR estimate from 25% to 18% through FY2028.'),
        ('Pillar 3 — Path to Superior Profitability: DELAYED, NOT ABANDONED.',
         'Management is making a deliberate, sizeable investment that delays profitability recovery. We now expect Non-GAAP operating margin to trough in H1 2026 (~17-19%) before recovering to 22-25% by FY2027-2028 as the First Party Brand generates operating leverage. EPS recovery begins in FY2027E as investment normalises.'),
    ]

    for bold_text, body_text in thesis_points:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(bold_text + ' ')
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(10)
        r1.font.color.rgb = NAVY
        r2 = p.add_run(body_text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10)

    add_heading(doc, '3.2 Key Risks (Updated)', level=2)

    risks = [
        ('Regulatory escalation (HIGH):', 'Temu faces ongoing investigations in the US (CPSC product safety), EU (Digital Services Act compliance), and ongoing CNIPA actions in China. Any material fines or platform restrictions could be a significant earnings headwind.'),
        ('Investment cycle overshoot (MEDIUM):', 'If the RMB100B commitment is front-loaded or materially exceeded, near-term EPS suppression could persist well into FY2027.'),
        ('Online marketing revenue stagnation (MEDIUM):', 'The 2.5% YoY growth in marketing services suggests Pinduoduo\'s domestic advertising revenue may be approaching maturity. Sustained stagnation would cap long-term revenue growth.'),
        ('Tariff and trade policy (MEDIUM):', 'US Section 301 tariffs and the expiration of the de minimis exemption for low-value parcels create structural headwinds for Temu\'s original marketplace model.'),
        ('Macro China softness (LOW-MEDIUM):', 'Consumer sentiment in China remains fragile. A deeper-than-expected domestic consumption slowdown could impact Pinduoduo GMV growth.'),
    ]

    for prefix, body in risks:
        add_bullet(doc, body, bold_prefix=prefix)

    # ══════════════════════════════════════════════════════════════
    # PAGE 8–10: VALUATION & ESTIMATES
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, '4. VALUATION & UPDATED ESTIMATES', level=1)

    add_heading(doc, '4.1 Estimate Revisions', level=2)
    add_body(doc, 'We reduce our revenue and earnings estimates meaningfully to reflect: (a) the confirmed Q1 2026 EPS run-rate; (b) management\'s qualitative guidance on continued near-term investment; and (c) our reduced Temu revenue growth assumptions through FY2028. The table below summarises our revised estimates.')

    est_data = [
        ['', 'FY2025E', 'FY2026E (Prior)', 'FY2026E (New)', 'Δ%', 'FY2027E (Prior)', 'FY2027E (New)', 'Δ%'],
        ['Revenue (RMB B)', '421', '450', '432', '−4.0%', '510', '498', '−2.4%'],
        ['YoY Growth', '+10.5%', '+7.0%', '+2.6%', '—', '+13.3%', '+15.3%', '—'],
        ['Gross Margin', '56.1%', '56.5%', '55.5%', '−100bps', '57.0%', '56.5%', '−50bps'],
        ['Non-GAAP Op. Margin', '18.3%', '21.5%', '18.8%', '−270bps', '23.0%', '21.5%', '−150bps'],
        ['Non-GAAP Net Income (RMB B)', '48.2', '67.5', '45.0', '−33.3%', '84.0', '59.0', '−29.8%'],
        ['Non-GAAP EPS/ADS (RMB)', '46.8', '65.5', '43.7', '−33.3%', '81.5', '57.3', '−29.7%'],
        ['Non-GAAP EPS/ADS (US$)', '$6.45', '$9.02', '$6.02', '−33.3%', '$11.23', '$7.89', '−29.7%'],
    ]

    tbl3 = doc.add_table(rows=len(est_data), cols=8)
    tbl3.style = 'Table Grid'
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    est_widths = [Inches(1.6), Inches(0.8), Inches(0.85), Inches(0.85), Inches(0.5), Inches(0.85), Inches(0.85), Inches(0.5)]

    for i, row_data in enumerate(est_data):
        for j, val in enumerate(row_data):
            cell = tbl3.cell(i, j)
            try:
                tbl3.columns[j].width = est_widths[j]
            except Exception:
                pass
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(8)
            if i == 0:
                r.bold = True
                r.font.color.rgb = WHITE
                set_cell_bg(cell, '1B4F8A')
            elif i % 2 == 0:
                set_cell_bg(cell, 'EBF5FB')
            if j == 0 and i > 0:
                r.bold = True
            if '−33' in val or '−29' in val or '−27' in val or '−4.0' in val:
                r.font.color.rgb = RED

    add_chart(doc, 'chart_08_estimate_revisions.png', width=5.8,
              caption='Figure 8: Revenue & EPS Estimate Revisions — Prior vs. New Estimates\nSource: Analyst estimates; PDD Holdings Q1 2026 Earnings Release (May 27, 2026)')

    add_heading(doc, '4.2 Revised Price Target: $125 (From $165)', level=2)
    add_body(doc, 'We revise our 12-month price target from $165 to $125, reflecting our reduced earnings estimates and a modest compression in the appropriate target multiple given near-term earnings uncertainty. Our price target is derived using a blended methodology:')

    pt_data = [
        ['Methodology', 'FY2026E / FY2027E Basis', 'Multiple / Rate', 'Implied Value', 'Weight', 'Contribution'],
        ['P/E (FY2027E Non-GAAP)', 'US$7.89/ADS', '16.0x', '$126', '40%', '$50'],
        ['EV/EBITDA (FY2027E)', 'US$10.2B EBITDA', '12.0x EV/EBITDA', '$118', '30%', '$35'],
        ['DCF (10yr, WACC 10.5%)', '8.5% terminal growth', 'g = 3%', '$145', '20%', '$29'],
        ['Sum-of-Parts (PDD+Temu)', 'Segment blended', 'PDD 12x + Temu 8x', '$138', '10%', '$14'],
        ['Blended Target', '', '', '', '100%', '$128'],
        ['Rounded Price Target', '', '', '', '', '$125'],
    ]

    tbl4 = doc.add_table(rows=len(pt_data), cols=6)
    tbl4.style = 'Table Grid'
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(pt_data):
        for j, val in enumerate(row_data):
            cell = tbl4.cell(i, j)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(8.5)
            if i == 0:
                r.bold = True
                r.font.color.rgb = WHITE
                set_cell_bg(cell, '1B4F8A')
            elif i >= len(pt_data) - 2:
                set_cell_bg(cell, 'FFF3CD')
                r.bold = True
                r.font.color.rgb = NAVY
            elif i % 2 == 0:
                set_cell_bg(cell, 'EBF5FB')
            if j == 0 and i > 0:
                r.bold = True

    add_body(doc, 'At our revised price target of $125, PDD would trade at 15.8x FY2027E Non-GAAP EPS — a meaningful discount to Amazon (35x P/E) and comparable to Sea Limited (22x), appropriate given the near-term investment cycle and regulatory uncertainty. With the stock at ~$86 post-earnings, the implied upside to our base case PT is +45%, which comfortably supports our BUY rating.', space_after=6)

    add_chart(doc, 'chart_09_peer_valuation.png', width=5.8,
              caption='Figure 9: Forward P/E & EV/EBITDA Peer Comparison — Post-Earnings\nSource: Bloomberg; analyst estimates (May 2026); PDD Holdings Q1 2026 Earnings Release')
    add_chart(doc, 'chart_10_price_target_scenarios.png', width=5.5,
              caption='Figure 10: Price Target Scenario Analysis — Bull / Base / Bear\nSource: Analyst estimates; PDD Holdings Q1 2026 Earnings Release (May 27, 2026)')

    # ══════════════════════════════════════════════════════════════
    # PAGE 11-12: SOURCES & DISCLOSURES
    # ══════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'SOURCES & REFERENCES', level=1)

    add_heading(doc, 'Primary Sources — Q1 2026 Earnings Materials', level=2)

    sources = [
        ('PDD Holdings Q1 2026 Earnings Release (May 27, 2026)',
         'https://investor.pddholdings.com/',
         'Press release announcing unaudited Q1 2026 financial results. Primary source for all revenue, profit, and EPS figures cited in this report.'),
        ('PDD Holdings Form 6-K (Filed May 27, 2026)',
         'https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001737806&type=6-K&dateb=&owner=include&count=40',
         'SEC EDGAR filing confirming Q1 2026 unaudited financial results. Accession: 0001104659-26-034813.'),
        ('PDD Holdings Q1 2026 Earnings Call Transcript (May 27, 2026)',
         'https://www.benzinga.com/insights/news/26/05/52804676/pdd-holdings-q1-2026-earnings-call-complete-transcript',
         'Complete transcript of management prepared remarks and analyst Q&A session. Source for all management quotations in this report.'),
        ('Bloomberg Consensus Estimates (as of May 26, 2026)',
         'https://www.bloomberg.com/quote/PDD:US',
         'Consensus revenue estimate US$15.94B; consensus Non-GAAP EPS US$2.13 per ADS. Used as the basis for all beat/miss calculations.'),
    ]

    for i, (title, url, desc) in enumerate(sources):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(3)
        r_num = p.add_run(f'[{i+1}]  ')
        r_num.bold = True
        r_num.font.name = 'Times New Roman'
        r_num.font.size = Pt(9)
        r_num.font.color.rgb = NAVY
        add_hyperlink(p, title, url)
        r_desc = p.add_run(f' — {desc}')
        r_desc.font.name = 'Times New Roman'
        r_desc.font.size = Pt(9)
        r_desc.font.color.rgb = GRAY

    add_heading(doc, 'Supporting References', level=2)

    supp = [
        ('StockTitan — PDD Q1 2026 Earnings Summary',
         'https://www.stocktitan.net/news/PDD/pdd-holdings-announces-first-quarter-2026-unaudited-financial-2iwyuhtjni8p.html'),
        ('PDD Holdings Investor Relations',
         'https://investor.pddholdings.com/'),
        ('PDD Holdings SEC EDGAR Filings',
         'https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=0001737806&type=&dateb=&owner=include&count=40&search_text='),
    ]

    for title, url in supp:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(2)
        r_bullet = p.add_run('•  ')
        r_bullet.font.name = 'Times New Roman'
        r_bullet.font.size = Pt(9)
        r_bullet.font.color.rgb = NAVY
        add_hyperlink(p, title, url)

    divider(doc)

    add_heading(doc, 'ANALYST CERTIFICATION & DISCLOSURES', level=2)
    disc_text = (
        'The views expressed in this report accurately reflect the personal views of the analyst(s) responsible for this report about the subject securities '
        'and/or issuers referenced herein. No part of the analyst\'s compensation was, is, or will be, directly or indirectly, related to the specific '
        'recommendations or views expressed in this report. This report is produced for informational purposes only and does not constitute an offer or '
        'solicitation to buy or sell securities. Past performance is not indicative of future results. Investors should conduct their own due diligence. '
        'This earnings update was prepared within 24 hours of the earnings release on May 27, 2026, consistent with institutional research standards.'
    )
    p_disc = doc.add_paragraph()
    p_disc.paragraph_format.left_indent = Inches(0.1)
    r_disc = p_disc.add_run(disc_text)
    r_disc.font.name = 'Times New Roman'
    r_disc.font.size = Pt(8)
    r_disc.font.color.rgb = GRAY
    r_disc.italic = True

    # ── Footer on all pages ───────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        ft_para = footer.paragraphs[0]
        ft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ft_run = ft_para.add_run(
            'PDD Holdings Inc. (NASDAQ: PDD) — Q1 2026 Earnings Update  |  May 27, 2026  |  '
            'Rating: BUY  |  Price Target: US$125  |  CONFIDENTIAL — FOR PROFESSIONAL INVESTORS ONLY'
        )
        ft_run.font.name = 'Times New Roman'
        ft_run.font.size = Pt(7.5)
        ft_run.font.color.rgb = GRAY

    doc.save(OUT_FILE)
    size_kb = os.path.getsize(OUT_FILE) / 1024
    print(f"\n  ✓  {OUT_FILE}  ({size_kb:.0f} KB)")


if __name__ == '__main__':
    print("Building PDD Holdings Q1 2026 Earnings Update Report...")
    build()
    print("Done.")
